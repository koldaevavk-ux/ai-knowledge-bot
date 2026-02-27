import os
import fitz  # PyMuPDF
from docx import Document
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
from app.config import config

class DocumentLoader:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def load_pdf(self, file_path: str) -> str:
        """Загрузка текста из PDF"""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    
    def load_docx(self, file_path: str) -> str:
        """Загрузка текста из DOCX"""
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    
    def load_all_documents(self) -> List[LangChainDocument]:
        """Загрузка всех документов из папки"""
        documents = []
        
        for filename in os.listdir(config.DOCUMENTS_PATH):
            file_path = os.path.join(config.DOCUMENTS_PATH, filename)
            
            try:
                if filename.endswith('.pdf'):
                    text = self.load_pdf(file_path)
                elif filename.endswith('.docx'):
                    text = self.load_docx(file_path)
                else:
                    continue
                
                # Разбиваем на чанки
                chunks = self.text_splitter.split_text(text)
                
                # Создаем документы с метаданными
                for i, chunk in enumerate(chunks):
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": filename,
                            "chunk": i
                        }
                    )
                    documents.append(doc)
                
                print(f"✅ Загружен: {filename} ({len(chunks)} чанков)")
            
            except Exception as e:
                print(f"❌ Ошибка загрузки {filename}: {e}")
        
        return documents
